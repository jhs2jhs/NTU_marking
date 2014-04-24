from docx import *
from docx.shared import Inches
import pprint
import codecs
import sys
import smtplib
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.header import Header
from email import encoders
import CONFIG_test as CONFIG

doc_level = 0
current = 'pseudo_code'
students = {}
student = {}

def get_grade(s):
	if s >= 70:
		return "First"
	if s >= 67 and s < 70:
		return "High 2:1"
	if s >= 63 and s < 67:
		return "Mid 2:1"
	if s >= 60 and s < 64:
		return "Low 2:1"
	if s >= 57 and s < 60:
		return "High 2:2"
	if s >= 53 and s < 57:
		return "Mid 2:2"
	if s >= 50 and s < 54:
		return "Low 2:2"
	if s >= 47 and s < 50:
		return "High 3rd"
	if s >= 43 and s < 47:
		return "Mid 3rd"
	if s >= 40 and s < 44:
		return "Low 3rd"
	if s >= 35 and s < 40:
		return "Marginal Fail"
	if s < 35:
		return "Fail"

def parse_student (t, post):
	global doc_level, student, current
	if t == '' and not post == '':
		doc_level = 0
		## save previous student into students
		if not student.has_key('student_id'):
			#print "student_id empty"
			#sys.exit(1)
			print 
		else:
			students[student['student_id']] = student
			#pprint.pprint(student)
			#print '\n\n\n'
		## parse name and student id
		name_id = post
		name_id = name_id.split('(')
		studnet_name = name_id[0].strip()
		student_id = name_id[1].replace(')', '').strip()
		student = {'parts':{}, 'student_name': studnet_name, 'student_id': student_id}
	if 'part a:' in t.lower():
		score = t.lower().split(':')[1]
		doc_level = 1
		student['parts'][doc_level] = {'score':score, 'pseudo_code':[], 'js_code':[]}
	elif 'part b:' in t.lower():
		score = t.lower().split(':')[1]
		doc_level = 2
		student['parts'][doc_level] = {'score':score, 'pseudo_code':[], 'js_code':[]}
	elif 'part c:' in t.lower():
		score = t.lower().split(':')[1]
		doc_level = 3
		student['parts'][doc_level] = {'score':score, 'pseudo_code':[], 'js_code':[]}
	elif 'overall:' in t.lower():
		score = t.lower().split(':')
		doc_level = 4
		student['parts'][doc_level] = {'score':0, 'pseudo_code':[], 'js_code':[]}
	elif 'pseudo code:' in t.lower():
		current = 'pseudo_code'
	elif 'js code:' in t.lower():
		current = 'js_code'
	elif 'all:' in t.lower() or 'you have therefore' in t.lower():
		return
	else:
		#print t
		if doc_level == 0:
			return
		if t == '':
			return
		student['parts'][doc_level][current].append(t)


def calcualte_score_grade():
	highest = 0
	lowest = 0
	average = 0
	grades = {}
	for student_id in students:
		parts = students[student_id]['parts']
		score_total = 0
		for part_id in parts:
			part_score = parts[part_id]['score']
			part_score = int(part_score)
			score_total = score_total + part_score
		grade = get_grade(score_total)
		students[student_id]['student_score'] = score_total
		students[student_id]['student_grade'] = grade
		## stastic
		if score_total > highest:
			highest = score_total
		if score_total < lowest:
			lowest = score_total
		average = average + score_total
		if not grades.has_key(grade):
			grades[grade] = {}
		grades[grade][student_id] = score_total
	average = average * 1.0 / len(students)
	print "==== student performance stastics ===="
	print "Students counts:", len(students)
	print "Highest:", highest, 'Lowest:', lowest, 'Average:', average
	for grade in grades:
		print "Grade:", grade, 'has student counts:', len(grades[grade])


def parse_docx(doc_name):
	pre = ''
	document = Document(doc_name)
	for post in document.paragraphs:
		t = post.text.strip()  ## curent line text
		parse_student(pre, t)
		pre = t
	students[student['student_id']] = student
	pprint.pprint(students)
	print "Student counts:", len(students)
	calcualte_score_grade()



##########################################################################

def print_student_grade_only(file_name):
	f = codecs.open(file_name, 'w', encoding='utf-8')
	f.write(u'%s\t%s\t%s\t%s\t\n'%('student_id', 'student_name', 'student_score', 'student_grade'))
	for student_id in students:
		student = students[student_id]
		student_name = student['student_name']
		student_score = student['student_score']
		student_grade = student['student_grade']
		f.write(u'%s\t%s\t%s\t%s\t\n'%(student_id, student_name, student_score, student_grade))
	f.close()
	print "== print_student_grade_only: %s =="%(file_name)

def print_student_comments_in_one_file(file_name, title, show_score):
	doc = Document()
	for student_id in students:
		student = students[student_id]
		student_name = student['student_name']
		student_score = student['student_score']
		student_grade = student['student_grade']
		doc.add_heading(title, 0)
		doc.add_heading(u'Student ID: %s, Student Name: %s'%(student_id, student_name), 1)
		doc.add_heading(u'Final Grade: %s'%(student_grade), 1)
		parts = student['parts']
		for part_id in parts:
			part_label = u'Part-%s'%(part_id)
			part_score = parts[part_id]['score']
			if (part_id == 4):
				part_label = 'Overall'
				part_score = student_score
			if show_score == True:
				doc.add_heading(u'%s: %s :'%(part_label, part_score), 2)
			else:
				doc.add_heading(u'%s: '%(part_label), 2)
			doc.add_heading('Pseudo code', 4)
			for c in parts[part_id]['pseudo_code']:
				doc.add_paragraph(c, style='ListBullet2')
			doc.add_heading('JS code', 4)
			for c in parts[part_id]['js_code']:
				doc.add_paragraph(c, style='ListBullet2')
		doc.add_page_break()
	doc.save(file_name)
	print "== print_student_comments_in_one_file: %s =="%(file_name)

def print_student_comments_in_seperate_file(file_name_prefix, title, show_score):
	for student_id in students:
		student = students[student_id]
		student_name = student['student_name']
		student_score = student['student_score']
		student_grade = student['student_grade']
		doc = Document()
		doc.add_heading(title, 0)
		doc.add_heading(u'Student ID: %s, Student Name: %s'%(student_id, student_name), 1)
		doc.add_heading(u'Final Grade: %s'%(student_grade), 1)
		parts = student['parts']
		for part_id in parts:
			part_label = u'Part-%s'%(part_id)
			part_score = parts[part_id]['score']
			if (part_id == 4):
				part_label = 'Overall'
				part_score = student_score
			if show_score == True:
				doc.add_heading(u'%s: %s :'%(part_label, part_score), 2)
			else:
				doc.add_heading(u'%s: '%(part_label), 2)
			doc.add_heading('Pseudo code', 4)
			for c in parts[part_id]['pseudo_code']:
				doc.add_paragraph(c, style='ListBullet2')
			doc.add_heading('JS code', 4)
			for c in parts[part_id]['js_code']:
				doc.add_paragraph(c, style='ListBullet2')
		doc.add_page_break()
		file_name = u'%s%s.docx'%(file_name_prefix, student_id)
		doc.save(file_name)
		print "== print_student_comments_in_seperate_file: %s =="%(file_name)


def print_student_comments_to_email(file_name_prefix, email_title_prefix, is_test, mail_type):
	for student_id in students:
		student = students[student_id]
		student_name = student['student_name']
		student_score = student['student_score']
		student_grade = student['student_grade']
		parts = student['parts']
		parts_html = ''
		for part_id in parts:
			part_label = u'Part-%s'%(part_id)
			if (part_id == 4):
				part_label = 'Overall'
				part_score = student_score
			part_html_pseudo_code = ''
			for c in parts[part_id]['pseudo_code']:
				part_html_pseudo_code = u'%s<li>%s</li>'%(part_html_pseudo_code, c)
			part_html_pseudo_code = u'<ul>%s</ul>'%(part_html_pseudo_code)
			part_html_js_code = ''
			for c in parts[part_id]['js_code']:
				part_html_js_code = u'%s<li>%s</li>'%(part_html_js_code, c)
			part_html_js_code = u'<ul>%s</ul>'%(part_html_js_code)
			part_html = part_template%(part_label, part_html_pseudo_code, part_html_js_code)
			parts_html = u'%s<br>%s'%(parts_html, part_html)
		file_name = u'%s%s.docx'%(file_name_prefix, student_id)
		email_body = email_template%(student_name, student_id, student_grade, parts_html)
		email_title = u'Web-Based-Programing In-Class test results: %s'%(student_id)
		#print email_body
		if is_test == True:
			address_to = CONFIG.address_to_test
			send_email(mail_type, email_title, email_body, address_to)
			print "== print_student_comments_to_email: %s =="%(file_name)
			print "only in test"
			return
		else:
			address_to = '%s@ntu.ac.uk'%(student_id)
			send_email(mail_type, email_title, email_body, address_to)
			print "== print_student_comments_to_email: %s =="%(file_name)

part_template = u'''
<b>%s</b>
<ul>Pseudo Code: 
	<ul>%s</ul>
</ul>
<ul>JS Code: 
	<ul>%s</ul>
</ul>
'''
email_template = u'''
To <b>%s</b>: <br>
<br>
Your student_id is <b>%s</b>, and your final grade is <b>%s</b> <br>
<br>
Bellow are your comments details: <br>
%s
'''
def send_email(mail_type, email_title, email_body, address_to):
	#mail_type = 'gmail' # or ntu
	is_attachment = False
	if mail_type == 'gmail':
		smtp_server = 'smtp.gmail.com'
		smtp_port = 587
		address_from = CONFIG.gmail_address_from
		username = CONFIG.gmail_username
		password = CONFIG.gmail_password
	if mail_type == 'ntu':
		smtp_server = 'smtphost.ntu.ac.uk'
		smtp_port = 25
		address_from = CONFIG.ntu_address_from
		username = CONFIG.ntu_username
		password = CONFIG.ntu_password
	address_cc = CONFIG.address_cc
	msg = MIMEMultipart('alternative')
	msg.set_charset('utf-8')
	msg['FROM'] = address_from
	msg['TO'] = address_to
	msg['CC'] = address_cc
	msg['Subject'] = email_title
	part2 = MIMEText(email_body, 'html', 'utf-8')
	msg.attach(part2)
	if is_attachment == True:
		### this is for attachment, disabled at moment
		#part = MIMEBase('application', "octet-stream")
		#part.set_payload(open('comments/H-N0491912.docx',"rb").read() )
		#encoders.encode_base64(part)
		#part.add_header('Content-Disposition', 'attachment; filename="H-N0491912.docx"')
		#msg.attach(part)
		#print msg
		print 'no attachment now'
	server = smtplib.SMTP(smtp_server, smtp_port)
	server.ehlo()
	server.starttls()
	server.login(username, password)
	server.sendmail(address_from, address_to, msg.as_string())
	server.quit()


def main_in_group(doc_name, group_name):
	doc_level = 0
	current = 'pseudo_code'
	student = {}
	parse_docx(doc_name)
	print_student_grade_only('./output_files/'+group_name+'-grade.txt')
	print_student_comments_in_one_file('./output_files/'+group_name+'-comments_in_one_file_with_score.docx', 'Web-Based-Programing In-Class test result', True)
	print_student_comments_in_one_file('./output_files/'+group_name+'-comments_in_one_file.docx', 'Web-Based-Programing In-Class test result', False)
	print_student_comments_in_seperate_file('./comments/'+group_name+'-', 'Web-Based-Programing In-Class test result', False)
	### ready to send email
	send_email_in_test = True   ############### make sure you changed here ##################
	mail_type = 'gmail' # or ntu ############### make sure you changed here ##################
	print_student_comments_to_email('./comments/'+group_name+'-', 'Web-Based-Programing In-Class test result', send_email_in_test, mail_type)


if __name__ == "__main__":
	main_in_group('./input_files/Marking-Group-H.docx', 'H') ############### make sure you changed here ##################
	main_in_group('./input_files/Marking-Group-F.docx', 'F') ############### make sure you changed here ##################





